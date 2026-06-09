# backend/app/batch_word/router.py
"""
批量转 Word - API 路由
端点：上传模板 / AI标注 / 保存映射 / 在线预览 / 批量生成 / 下载
"""
import io
import uuid
from datetime import datetime, timezone
from pathlib import Path

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from fastapi.responses import StreamingResponse, FileResponse

from sqlalchemy.ext.asyncio import AsyncSession

from ..auth.models import User
from ..core.database import async_session_maker, get_db
from ..core.dependencies import get_current_user
from ..core.quota import QuotaGuard
from ..core.usage_service import increment_usage
from ..files import ugc_registry_service as ugc_registry
from ..utils.logger import get_logger
from . import schemas as S
from . import service
from .auto_annotator import auto_annotate

logger = get_logger("batch_word.router")

router = APIRouter(prefix="/api/batch-word", tags=["batch-word"])

_BW_PROJECT_ROOT = Path(__file__).resolve().parents[3]


_BANNED_ANNOTATE_TEXT = (
    "准考证",
    "演示",
    "说明",
    "Q文档",
    "http://",
    "https://",
    "www.",
)


def _is_unsafe_original_text(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if len(t) <= 1:
        return True
    if t in ("姓名", "身份证号", "考试时间", "考试地点", "照片", "头像", "编号", "序号"):
        return True
    return any(k in t for k in _BANNED_ANNOTATE_TEXT)


def _collect_doc_full_text(doc) -> str:
    """收集 doc 全文（正文 + 表格 + 页眉页脚）"""
    parts = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = (p.text or "").strip()
                    if text:
                        parts.append(text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)
        for p in section.footer.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)
    return "\n".join(parts)


def _guess_image_column(excel_columns: list[str]) -> str | None:
    """推断图片字段列名"""
    keys = ("照片", "头像", "图片", "证件照", "image", "photo", "avatar", "img")
    for c in excel_columns or []:
        lc = str(c).lower()
        if any(k in lc for k in keys):
            return c
    return None


def _normalize_placeholder(raw_placeholder: str, fallback_column: str = "") -> str:
    """规范化占位符，统一为 {字段名}。"""
    text = str(raw_placeholder or "").strip()
    fallback = str(fallback_column or "").strip()
    if not text and fallback:
        text = fallback
    if not text:
        return ""
    import re
    name = re.sub(r"^\{+|\}+$", "", text).strip()
    if not name and fallback:
        name = re.sub(r"^\{+|\}+$", "", fallback).strip()
    if not name:
        return ""
    return f"{{{name}}}"


# ==================== 上传模板 ====================

@router.post("/upload-template", response_model=S.UploadTemplateResponse)
async def upload_template(
    file: UploadFile = File(...),
    _user: User = Depends(get_current_user),
):
    """上传 Word 模板(.docx)，解析标注符并返回摘要"""
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "仅支持 .docx 格式")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(400, "文件不能超过 50MB")

    template_id, path = service.save_template(content, file.filename)
    placeholders = service.extract_placeholders(path)
    summary = service.extract_text_summary(path)
    has_images = service.has_images(path)
    cfg = service.load_template_config(template_id)
    saved_mappings = cfg.get("mappings", [])
    saved_filename_pattern = cfg.get("filename_pattern", "文档_{_index}")
    has_saved_config = bool(cfg.get("has_saved_config"))
    saved_editor_html = str(cfg.get("editor_html") or "")

    return S.UploadTemplateResponse(
        template_id=template_id,
        placeholders=placeholders,
        text_summary=summary,
        has_images=has_images,
        has_saved_config=has_saved_config,
        saved_mappings=[S.MappingItem(**m) for m in saved_mappings],
        saved_filename_pattern=saved_filename_pattern,
        saved_editor_html=saved_editor_html,
    )


@router.get("/template/{template_id}", response_model=S.UploadTemplateResponse)
async def get_template_info(
    template_id: str,
    _user: User = Depends(get_current_user),
):
    """读取已上传模板信息（用于历史记录“修改”回填）"""
    try:
        path = service.get_template_path(template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    placeholders = service.extract_placeholders(path)
    summary = service.extract_text_summary(path)
    has_images = service.has_images(path)
    cfg = service.load_template_config(template_id)
    saved_mappings = cfg.get("mappings", [])
    saved_filename_pattern = cfg.get("filename_pattern", "文档_{_index}")
    has_saved_config = bool(cfg.get("has_saved_config"))
    saved_editor_html = str(cfg.get("editor_html") or "")
    return S.UploadTemplateResponse(
        template_id=template_id,
        placeholders=placeholders,
        text_summary=summary,
        has_images=has_images,
        has_saved_config=has_saved_config,
        saved_mappings=[S.MappingItem(**m) for m in saved_mappings],
        saved_filename_pattern=saved_filename_pattern,
        saved_editor_html=saved_editor_html,
    )


# ==================== AI 自动标注 ====================

@router.post("/auto-annotate", response_model=S.AutoAnnotateResponse)
async def auto_annotate_template(
    req: S.AutoAnnotateRequest,
    _user: User = Depends(get_current_user),
):
    """LLM 分析 Word 全文，返回标注映射建议"""
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    # 获取更完整的文本用于 LLM 分析
    from docx import Document
    doc = Document(str(path))
    full_text = _collect_doc_full_text(doc) or service.extract_text_summary(path)
    has_image_slot = len(doc.inline_shapes) > 0

    suggestions = await auto_annotate(
        doc_text=full_text,
        excel_columns=req.excel_columns,
        has_image_slot=has_image_slot,
    )

    # 构建标注后的文本预览
    preview = full_text[:1000]
    for s in suggestions:
        if s.get("original_text") and s.get("placeholder"):
            preview = preview.replace(
                s["original_text"], s["placeholder"], 1
            )

    # 自动保存 AI 建议，便于前端刷新后继续编辑
    auto_mappings = []
    for s in suggestions:
        auto_mappings.append({
            "placeholder": s.get("placeholder", ""),
            "column": s.get("column", ""),
            "type": "image" if any(x in s.get("column", "") for x in ("照片", "头像", "图片", "image")) else "text",
        })
    if auto_mappings:
        service.save_template_mappings(req.template_id, auto_mappings)

    return S.AutoAnnotateResponse(
        suggestions=[
            S.AnnotationSuggestion(**s) for s in suggestions
        ],
        annotated_text_preview=preview[:1000],
    )


@router.post("/save-mappings", response_model=S.SaveMappingsResponse)
async def save_mappings(
    req: S.SaveMappingsRequest,
    _user: User = Depends(get_current_user),
):
    """保存映射（前端微调自动保存）"""
    try:
        _ = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")
    mappings = [m.model_dump() for m in req.mappings]
    service.save_template_mappings(
        req.template_id,
        mappings,
        filename_pattern=req.filename_pattern,
        editor_html=req.editor_html,
    )
    if req.record_history:
        template_path = service.get_template_path(req.template_id)
        service.append_history_record({
            "user_id": str(_user.id),
            "task_id": uuid.uuid4().hex[:12],
            "template_id": req.template_id,
            "template_file_name": (req.template_file_name or template_path.name or "").strip(),
            "source_file_id": req.source_file_id or "",
            "source_file_name": req.source_file_name or "",
            "filename_pattern": req.filename_pattern or "文档_{_index}",
            "total": 0,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "download_url": "",
        })
    cfg = service.load_template_config(req.template_id)
    return S.SaveMappingsResponse(
        ok=True,
        saved_count=len(mappings),
        saved_filename_pattern=str(cfg.get("filename_pattern", "文档_{_index}")),
    )


# ==================== AI 一键标注（修改 docx 本体） ====================

@router.post("/ai-annotate-doc", response_model=S.AIAnnotateDocResponse)
async def ai_annotate_doc(
    req: S.AIAnnotateDocRequest,
    _user: User = Depends(get_current_user),
):
    """
    AI 一键标注：LLM 分析 Word 全文，直接在 docx 中将 demo 数据替换为 {字段}，
    覆盖保存后返回高保真 HTML + 映射列表。
    """
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    from docx import Document
    doc = Document(str(path))
    full_text = _collect_doc_full_text(doc) or service.extract_text_summary(path)
    has_image_slot = len(doc.inline_shapes) > 0

    suggestions = await auto_annotate(
        doc_text=full_text,
        excel_columns=req.excel_columns,
        has_image_slot=has_image_slot,
    )

    # 规则优先：用 sample_row 值做精确匹配（比纯 LLM 更稳定）
    deterministic = []
    sample_row = req.sample_row or {}
    for col in req.excel_columns:
        raw_val = sample_row.get(col)
        val = str(raw_val).strip() if raw_val is not None else ""
        if not val or _is_unsafe_original_text(val):
            continue
        if val in full_text:
            deterministic.append({
                "original_text": val,
                "placeholder": f"{{{col}}}",
                "column": col,
                "confidence": 0.99,
            })

    # 合并：规则优先，LLM 作为补充
    merged = []
    seen_pair = set()
    for s in deterministic + suggestions:
        orig = (s.get("original_text") or "").strip()
        ph = _normalize_placeholder(s.get("placeholder"), s.get("column"))
        col = (s.get("column") or "").strip()
        if _is_unsafe_original_text(orig):
            continue
        if not (orig and ph and col):
            continue
        key = (orig, ph)
        if key in seen_pair:
            continue
        seen_pair.add(key)
        merged.append({
            "original_text": orig,
            "placeholder": ph,
            "column": col,
            "confidence": float(s.get("confidence", 0.8)),
        })

    image_col = _guess_image_column(req.excel_columns)
    # 构造替换映射 original_text -> {placeholder}
    replace_map: dict[str, str] = {}
    result_mappings: list[S.MappingItem] = []
    for s in merged:
        orig = (s.get("original_text") or "").strip()
        ph = (s.get("placeholder") or "").strip()
        col = (s.get("column") or "").strip()
        if orig and ph and col and orig not in replace_map:
            replace_map[orig] = ph
        if ph and col:
            is_img = any(
                k in col for k in ("照片", "头像", "图片", "image")
            )
            result_mappings.append(S.MappingItem(
                placeholder=ph,
                column=col,
                type="image" if is_img else "text",
            ))

    # 先执行文本替换
    html = service.annotate_doc_with_replacements(path, replace_map)

    # 图片位兜底：若文档有图片且映射中没有图片项，则将首个图片位标注为 {图片列}
    has_image_mapping = any(m.type == "image" for m in result_mappings)
    if has_image_slot and image_col and not has_image_mapping:
        html = service.manual_annotate_image_in_doc(path, image_col)
        result_mappings.append(S.MappingItem(
            placeholder=f"{{{image_col}}}",
            column=image_col,
            type="image",
        ))

    # 持久化映射
    service.save_template_mappings(
        req.template_id,
        [m.model_dump() for m in result_mappings],
    )

    return S.AIAnnotateDocResponse(html=html, mappings=result_mappings)


# ==================== 手工标注（修改 docx 本体） ====================

@router.post("/manual-annotate", response_model=S.ManualAnnotateResponse)
async def manual_annotate(
    req: S.ManualAnnotateRequest,
    _user: User = Depends(get_current_user),
):
    """
    手工标注：将选中的原始文本替换为 {field_name}，
    覆盖保存 docx 后返回新的高保真 HTML。
    """
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    if not req.original_text.strip() or not req.field_name.strip():
        raise HTTPException(400, "original_text 和 field_name 不能为空")

    html = service.manual_annotate_in_doc(
        path, req.original_text.strip(), req.field_name.strip()
    )
    return S.ManualAnnotateResponse(html=html)


@router.post("/manual-annotate-image", response_model=S.ManualAnnotateResponse)
async def manual_annotate_image(
    req: S.ManualAnnotateImageRequest,
    _user: User = Depends(get_current_user),
):
    """手工图片标注：右键图片后替换为字段占位符"""
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    if not req.field_name.strip():
        raise HTTPException(400, "field_name 不能为空")

    html = service.manual_annotate_image_in_doc(
        path, req.field_name.strip(), req.embed_id
    )
    return S.ManualAnnotateResponse(html=html)


@router.post("/restore-template", response_model=S.RestoreTemplateResponse)
async def restore_template(
    req: S.RestoreTemplateRequest,
    _user: User = Depends(get_current_user),
):
    """一键还原：恢复到上传时原始模板状态"""
    try:
        path = service.restore_template_to_original(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在或原始副本缺失")

    placeholders = service.extract_placeholders(path)
    summary = service.extract_text_summary(path)
    has_images = service.has_images(path)
    html = service.build_preview_html(
        template_path=path,
        mappings=[],
        row_data={},
        mode="annotated",
    )
    return S.RestoreTemplateResponse(
        template_id=req.template_id,
        placeholders=placeholders,
        text_summary=summary,
        has_images=has_images,
        html=html,
    )


@router.post("/preview-html", response_model=S.PreviewHtmlResponse)
async def preview_html(
    req: S.PreviewHtmlRequest,
    _user: User = Depends(get_current_user),
):
    """在线预览 HTML（用于核对标注准确性）"""
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    mappings = [m.model_dump() for m in req.mappings]
    html = service.build_preview_html(
        template_path=path,
        mappings=mappings,
        row_data=req.row_data or {},
        mode=req.mode or "filled",
    )
    return S.PreviewHtmlResponse(html=html)


# ==================== 批量生成 ====================

@router.post("/generate", response_model=S.GenerateResponse)
async def generate_batch(
    req: S.GenerateRequest,
    _user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    _quota=Depends(QuotaGuard("batch_word_monthly")),
):
    """批量生成填充后的 Word 文档并打包 ZIP"""
    await increment_usage(_user.id, "batch_word_count", db)
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    if not req.rows:
        raise HTTPException(400, "行数据不能为空")

    mappings = [m.model_dump() for m in req.mappings]

    task_id, zip_path = service.batch_generate(
        template_path=path,
        mappings=mappings,
        rows=req.rows,
        filename_pattern=req.filename_pattern,
    )

    removed = service.append_history_record({
        "user_id": str(_user.id),
        "task_id": task_id,
        "template_id": req.template_id,
        "template_file_name": req.template_file_name or "",
        "source_file_id": req.source_file_id or "",
        "source_file_name": req.source_file_name or "",
        "filename_pattern": req.filename_pattern,
        "total": len(req.rows),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "download_url": f"/api/batch-word/download/{task_id}",
    })

    try:
        zip_rel = str(zip_path.resolve().relative_to(_BW_PROJECT_ROOT.resolve()).as_posix())
    except ValueError:
        zip_rel = str(zip_path.as_posix())
    zip_sz = int(zip_path.stat().st_size) if zip_path.exists() else 0

    async with async_session_maker() as db_sess:
        try:
            for tid in removed:
                await ugc_registry.mark_batch_word_export_deleted(db_sess, tid)
            await ugc_registry.upsert_batch_word_export(
                db_sess,
                task_id=task_id,
                user_id=str(_user.id),
                template_id=req.template_id,
                template_file_name=req.template_file_name or "",
                source_file_id=req.source_file_id or "",
                filename_pattern=req.filename_pattern,
                zip_rel_path=zip_rel,
                total=len(req.rows),
                zip_size_bytes=zip_sz,
            )
            await db_sess.commit()
        except Exception as exc:
            logger.warning("batch_word_exports 注册表写入失败: %s", exc)
            await db_sess.rollback()

    return S.GenerateResponse(
        task_id=task_id,
        total=len(req.rows),
        download_url=f"/api/batch-word/download/{task_id}",
    )


# ==================== 预览（单行） ====================

@router.post("/preview")
async def preview_single(
    req: S.PreviewRequest,
    _user: User = Depends(get_current_user),
):
    """预览单条文档，返回 .docx 二进制流"""
    try:
        path = service.get_template_path(req.template_id)
    except FileNotFoundError:
        raise HTTPException(404, "模板不存在")

    mappings = [m.model_dump() for m in req.mappings]
    doc_bytes = service.fill_document(path, mappings, req.row_data)

    return StreamingResponse(
        io.BytesIO(doc_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=preview.docx"},
    )


# ==================== 下载 ZIP ====================

@router.get("/history", response_model=S.BatchWordHistoryResponse)
async def list_history(
    _user: User = Depends(get_current_user),
):
    """获取当前用户历史转换清单"""
    records = service.list_history_records(str(_user.id), limit=50)
    return S.BatchWordHistoryResponse(
        items=[S.BatchWordHistoryItem(**r) for r in records]
    )


@router.delete("/history/{task_id}", response_model=S.BatchWordHistoryDeleteResponse)
async def delete_history(
    task_id: str,
    _user: User = Depends(get_current_user),
):
    """删除历史记录，同时删除服务端模板文档与结果 zip"""
    record = service.get_history_record(str(_user.id), task_id)
    if not record:
        raise HTTPException(404, "历史记录不存在")

    template_id = str(record.get("template_id") or "").strip()
    if template_id:
        service.delete_template_artifacts(template_id, task_id=task_id)
    else:
        # 兼容无 template_id 的旧记录：仅删 zip
        service.delete_template_artifacts("", task_id=task_id)

    ok = service.remove_history_record(str(_user.id), task_id)
    if not ok:
        raise HTTPException(500, "删除历史记录失败")
    return S.BatchWordHistoryDeleteResponse(ok=True, deleted_task_id=task_id)


@router.get("/download/{task_id}")
async def download_zip(
    task_id: str,
    _user: User = Depends(get_current_user),
):
    """下载批量生成的 ZIP 包"""
    hit = service._find_file(f"{task_id}.zip")
    zip_path = hit if hit and hit.exists() else None
    if not zip_path:
        raise HTTPException(404, "文件不存在或已过期")

    return FileResponse(
        path=str(zip_path),
        media_type="application/zip",
        filename=f"batch_word_{task_id}.zip",
    )
