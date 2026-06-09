# backend/app/batch_word/service.py
"""
批量转 Word - 核心服务
模板解析 / 文档填充 / 图片替换 / ZIP 打包
"""
import io
import os
import re
import json
import uuid
import base64
import zipfile
import tempfile
from datetime import date, datetime
from html import escape
from urllib.request import urlopen
from typing import List, Dict, Any, Tuple
from pathlib import Path

from lxml import etree
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from ..utils.logger import get_logger

logger = get_logger("batch_word.service")

# ==================== 常量 ====================

PROJECT_ROOT = Path(__file__).resolve().parents[3]
UPLOAD_DIR = PROJECT_ROOT / "uploads" / "word_files"
PLACEHOLDER_RE = re.compile(r"\{([^}]+)\}")
MAX_SUMMARY_LEN = 500


def _normalize_placeholder(raw_placeholder: Any, fallback_column: Any = "") -> str:
    """
    规范化占位符，统一为 {字段名} 形式。
    - 支持输入: 姓名 / {姓名} / {{姓名}} / {{{姓名}}}
    - 若占位符为空，回退使用列名
    - 规范化失败返回空字符串
    """
    text = str(raw_placeholder or "").strip()
    fallback = str(fallback_column or "").strip()
    if not text and fallback:
        text = fallback
    if not text:
        return ""

    # 去除两侧任意层级花括号
    name = re.sub(r"^\{+|\}+$", "", text).strip()
    if not name and fallback:
        name = re.sub(r"^\{+|\}+$", "", fallback).strip()
    if not name:
        return ""
    return f"{{{name}}}"


def _normalize_nested_placeholder_braces(text: str) -> str:
    """
    归一化多层花括号占位符：
    - {{姓名}} / {{{姓名}}} -> {姓名}
    """
    if not text:
        return ""
    out = str(text)
    for _ in range(4):
        next_out = re.sub(r"\{\{+\s*([^{}]+?)\s*\}\}+", r"{\1}", out)
        if next_out == out:
            break
        out = next_out
    return out


def _cleanup_empty_brace_artifacts(text: str) -> str:
    """
    清理图片替换后残留的空花括号噪音：
    - {{}} / {{{}}} / {   } -> ""
    """
    if not text:
        return ""
    out = _normalize_nested_placeholder_braces(text)
    out = re.sub(r"\{+\s*\}+", "", out)
    return out


def _normalize_date_like_text(value: Any) -> str:
    """将日期样式值统一为 YYYY/M/D[ HH:MM[:SS]] 文本，非日期原样返回字符串。"""
    if value is None:
        return ""
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            if value.second:
                return f"{value.year}/{value.month}/{value.day} {value.hour:02d}:{value.minute:02d}:{value.second:02d}"
            return f"{value.year}/{value.month}/{value.day} {value.hour:02d}:{value.minute:02d}"
        return f"{value.year}/{value.month}/{value.day}"
    if isinstance(value, date):
        return f"{value.year}/{value.month}/{value.day}"

    text = str(value).strip()
    if not text:
        return ""
    text = re.sub(r'\\+"', '"', text)
    text = re.sub(r'^(["\'])+|(["\'])+$', "", text).strip()

    dt = re.search(
        r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})\s+(\d{1,2}):(\d{2})(?::(\d{2}))?",
        text,
    )
    if dt:
        y, mm, dd, hh, mi, ss = dt.groups()
        if ss:
            return f"{y}/{int(mm)}/{int(dd)} {int(hh):02d}:{int(mi):02d}:{int(ss):02d}"
        return f"{y}/{int(mm)}/{int(dd)} {int(hh):02d}:{int(mi):02d}"

    m = re.search(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", text)
    if m:
        return f"{m.group(1)}/{int(m.group(2))}/{int(m.group(3))}"
    zh = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", text)
    if zh:
        return f"{zh.group(1)}/{int(zh.group(2))}/{int(zh.group(3))}"
    return text


# ==================== 模板存储 ====================

def ensure_upload_dir() -> Path:
    """确保上传目录存在"""
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    return UPLOAD_DIR


def _today_dir() -> Path:
    d = UPLOAD_DIR / datetime.now().strftime("%Y-%m-%d")
    d.mkdir(parents=True, exist_ok=True)
    return d


def _find_file(pattern: str) -> Path | None:
    """在 UPLOAD_DIR 下递归查找匹配文件（跳过 history.json）。"""
    hits = list(UPLOAD_DIR.rglob(pattern))
    return hits[0] if hits else None


def save_template(file_bytes: bytes, filename: str) -> Tuple[str, Path]:
    """保存上传的模板文件到日期子目录，返回 (template_id, file_path)"""
    date_dir = _today_dir()
    ext = Path(filename).suffix or ".docx"
    template_id = uuid.uuid4().hex[:12]
    dest = date_dir / f"{template_id}{ext}"
    dest.write_bytes(file_bytes)
    original = date_dir / f"{template_id}.original{ext}"
    original.write_bytes(file_bytes)
    logger.info("模板已保存: id=%s path=%s", template_id, dest)
    return template_id, dest


def get_template_path(template_id: str) -> Path:
    """根据 template_id 递归查找模板路径"""
    for f in UPLOAD_DIR.rglob(f"{template_id}.*"):
        if ".original." not in f.name and ".mappings." not in f.name:
            return f
    raise FileNotFoundError(f"模板不存在: {template_id}")


def get_original_template_path(template_id: str) -> Path:
    """获取模板原始副本路径（上传后从未被修改的版本）"""
    hit = _find_file(f"{template_id}.original.*")
    if hit:
        return hit
    raise FileNotFoundError(f"模板原始副本不存在: {template_id}")


def restore_template_to_original(template_id: str) -> Path:
    """将当前模板恢复为上传时初始版本，返回模板路径"""
    current = get_template_path(template_id)
    original = get_original_template_path(template_id)
    current.write_bytes(original.read_bytes())
    return current


def _mapping_file_path(template_id: str) -> Path:
    """映射 sidecar 文件路径（与模板同目录）"""
    hit = _find_file(f"{template_id}.mappings.json")
    if hit:
        return hit
    # 新建时跟随模板所在目录，或回退到当日目录
    try:
        tpl = get_template_path(template_id)
        return tpl.parent / f"{template_id}.mappings.json"
    except FileNotFoundError:
        return _today_dir() / f"{template_id}.mappings.json"


def _history_file_path() -> Path:
    """历史记录文件路径（根级全局索引）"""
    ensure_upload_dir()
    return UPLOAD_DIR / "history.json"


def save_template_mappings(
    template_id: str,
    mappings: List[Dict[str, Any]],
    filename_pattern: str | None = None,
    editor_html: str | None = None,
) -> None:
    """保存模板映射（用于在线微调自动保存）"""
    # 去重：同一 placeholder 仅保留第一条，避免重复映射污染
    deduped: List[Dict[str, Any]] = []
    seen_placeholder = set()
    for m in mappings or []:
        if not isinstance(m, dict):
            continue
        ph = _normalize_placeholder(m.get("placeholder"), m.get("column"))
        col = str(m.get("column", "")).strip()
        typ = str(m.get("type", "text")).strip() or "text"
        if not ph or not col:
            continue
        key_ph = ph.lower()
        if key_ph in seen_placeholder:
            continue
        seen_placeholder.add(key_ph)
        deduped.append({
            "placeholder": ph,
            "column": col,
            "type": typ,
        })

    p = _mapping_file_path(template_id)
    old_filename_pattern = "文档_{_index}"
    old_editor_html = ""
    if p.exists():
        try:
            old_payload = json.loads(p.read_text(encoding="utf-8"))
            old_filename_pattern = str(
                old_payload.get("filename_pattern", "文档_{_index}")
            )
            old_editor_html = str(old_payload.get("editor_html") or "")
        except Exception:
            old_filename_pattern = "文档_{_index}"
            old_editor_html = ""
    final_filename_pattern = (
        filename_pattern.strip() if isinstance(filename_pattern, str) and filename_pattern.strip()
        else old_filename_pattern
    )
    final_editor_html = (
        editor_html if isinstance(editor_html, str) else old_editor_html
    )
    payload = {
        "template_id": template_id,
        "mappings": deduped,
        "filename_pattern": final_filename_pattern,
        "editor_html": final_editor_html,
    }
    p.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_template_mappings(template_id: str) -> List[Dict[str, Any]]:
    """读取模板映射"""
    cfg = load_template_config(template_id)
    return cfg.get("mappings", [])


def load_template_config(template_id: str) -> Dict[str, Any]:
    """读取模板配置（映射 + 文件名模式）"""
    p = _mapping_file_path(template_id)
    if not p.exists():
        return {
            "mappings": [],
            "filename_pattern": "文档_{_index}",
            "editor_html": "",
            "has_saved_config": False,
        }
    try:
        payload = json.loads(p.read_text(encoding="utf-8"))
        mappings = payload.get("mappings") or []
        if not isinstance(mappings, list):
            mappings = []
        filename_pattern = str(payload.get("filename_pattern") or "文档_{_index}")
        return {
            "mappings": mappings,
            "filename_pattern": filename_pattern,
            "editor_html": str(payload.get("editor_html") or ""),
            "has_saved_config": True,
        }
    except Exception:
        logger.warning("读取模板配置失败: %s", p)
        return {
            "mappings": [],
            "filename_pattern": "文档_{_index}",
            "editor_html": "",
            "has_saved_config": False,
        }


def append_history_record(record: Dict[str, Any]) -> List[str]:
    """追加历史记录（同用户+同模板仅保留最新一条）。
    返回因被替换而应从注册表标记删除的 task_id 列表。
    """
    removed_task_ids: List[str] = []
    p = _history_file_path()
    items: List[Dict[str, Any]] = []
    if p.exists():
        try:
            payload = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(payload, list):
                items = payload
        except Exception:
            logger.warning("读取历史记录失败: %s", p)
    user_id = str(record.get("user_id") or "").strip()
    template_id = str(record.get("template_id") or "").strip()
    template_file_name = str(record.get("template_file_name") or "").strip().lower()
    source_file_id = str(record.get("source_file_id") or "").strip()

    def is_same_template(item: Dict[str, Any]) -> bool:
        item_tpl_id = str(item.get("template_id") or "").strip()
        item_tpl_name = str(item.get("template_file_name") or "").strip().lower()
        item_source_id = str(item.get("source_file_id") or "").strip()

        # 1) template_id 完全一致，直接判定为同模板
        if template_id and item_tpl_id and item_tpl_id == template_id:
            return True

        # 2) 同名模板文件：同一个来源文件（或来源为空）视为同模板
        if template_file_name and item_tpl_name and item_tpl_name == template_file_name:
            if source_file_id and item_source_id:
                return item_source_id == source_file_id
            return True

        return False

    # 同一用户下，同一模板只保留一条最新记录，避免列表重复
    if user_id:
        next_items: List[Dict[str, Any]] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            same_user = str(item.get("user_id") or "").strip() == user_id
            same_tpl = is_same_template(item)
            if same_user and same_tpl:
                # 若新记录是“仅保存配置”（无下载地址），保留旧记录中的导出信息
                if not str(record.get("download_url") or "").strip():
                    if str(item.get("download_url") or "").strip():
                        record["download_url"] = str(item.get("download_url") or "")
                    if not str(record.get("task_id") or "").strip():
                        record["task_id"] = str(item.get("task_id") or "")
                    if not int(record.get("total") or 0):
                        record["total"] = int(item.get("total") or 0)
                old_task_id = str(item.get("task_id") or "").strip()
                should_cleanup_old_zip = bool(str(record.get("download_url") or "").strip())
                if old_task_id and should_cleanup_old_zip:
                    old_zip = _find_file(f"{old_task_id}.zip")
                    if old_zip and old_zip.exists():
                        old_zip.unlink(missing_ok=True)
                    removed_task_ids.append(old_task_id)
                continue
            next_items.append(item)
        items = next_items

    items.insert(0, record)
    p.write_text(json.dumps(items[:200], ensure_ascii=False, indent=2), encoding="utf-8")
    return removed_task_ids


def list_history_records(user_id: str, limit: int = 30) -> List[Dict[str, Any]]:
    """读取当前用户的历史记录"""
    p = _history_file_path()
    if not p.exists():
        return []
    try:
        payload = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        logger.warning("读取历史记录失败: %s", p)
        return []
    if not isinstance(payload, list):
        return []
    records = [r for r in payload if isinstance(r, dict) and r.get("user_id") == user_id]
    if not records:
        # 兼容旧版本：历史记录中无 user_id 时，回退显示公共记录
        records = [r for r in payload if isinstance(r, dict) and not r.get("user_id")]
    # 展示层去重：同一模板仅保留最新一条，兼容历史遗留重复记录
    deduped: List[Dict[str, Any]] = []
    seen_template = set()
    for item in records:
        template_id = str(item.get("template_id") or "").strip()
        template_file_name = str(item.get("template_file_name") or "").strip().lower()
        source_file_id = str(item.get("source_file_id") or "").strip()

        if template_file_name:
            dedupe_key = f"name:{template_file_name}|src:{source_file_id or '-'}"
        elif template_id:
            dedupe_key = f"id:{template_id}"
        else:
            dedupe_key = f"task:{str(item.get('task_id') or '').strip()}"

        if dedupe_key in seen_template:
            continue
        seen_template.add(dedupe_key)
        deduped.append(item)
    return deduped[:limit]


def remove_history_record(user_id: str, task_id: str) -> bool:
    """删除历史记录并返回是否删除成功"""
    p = _history_file_path()
    if not p.exists():
        return False
    try:
        payload = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        logger.warning("读取历史记录失败: %s", p)
        return False
    if not isinstance(payload, list):
        return False

    removed = False
    next_items: List[Dict[str, Any]] = []
    for item in payload:
        if not isinstance(item, dict):
            continue
        same_user = item.get("user_id") == user_id or not item.get("user_id")
        if same_user and item.get("task_id") == task_id and not removed:
            removed = True
            continue
        next_items.append(item)

    if removed:
        p.write_text(json.dumps(next_items, ensure_ascii=False, indent=2), encoding="utf-8")
    return removed


def get_history_record(user_id: str, task_id: str) -> Dict[str, Any] | None:
    """按 task_id 获取历史记录"""
    records = list_history_records(user_id, limit=500)
    for item in records:
        if item.get("task_id") == task_id:
            return item
    return None


def delete_template_artifacts(template_id: str, task_id: str | None = None) -> None:
    """
    删除模板相关服务端文件：
    - 模板 docx
    - mappings sidecar
    - 对应任务 zip（若提供 task_id）
    """
    if template_id:
        try:
            path = get_template_path(template_id)
            if path.exists():
                path.unlink(missing_ok=True)
        except Exception:
            pass

        mapping_path = _mapping_file_path(template_id)
        if mapping_path.exists():
            mapping_path.unlink(missing_ok=True)
        for original_path in UPLOAD_DIR.rglob(f"{template_id}.original.*"):
            if original_path.exists():
                original_path.unlink(missing_ok=True)

    if task_id:
        hit = _find_file(f"{task_id}.zip")
        if hit and hit.exists():
            hit.unlink(missing_ok=True)


# ==================== 模板解析 ====================

def _collect_paragraph_text(paragraph) -> str:
    """收集段落完整文本（合并所有 run）"""
    return "".join(run.text for run in paragraph.runs)


def _iter_all_paragraphs(doc: Document):
    """遍历文档所有段落（正文 + 表格 + 页眉页脚）"""
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

    for section in doc.sections:
        for para in section.header.paragraphs:
            yield para
        for para in section.footer.paragraphs:
            yield para


def extract_placeholders(template_path: Path) -> List[str]:
    """从模板中提取所有 {标注符}"""
    doc = Document(str(template_path))
    found = set()
    for para in _iter_all_paragraphs(doc):
        text = _collect_paragraph_text(para)
        for m in PLACEHOLDER_RE.finditer(text):
            found.add(m.group(0))
    return sorted(found)


def extract_text_summary(template_path: Path) -> str:
    """提取文档纯文本摘要"""
    doc = Document(str(template_path))
    parts = []
    total = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
            total += len(t)
            if total > MAX_SUMMARY_LEN:
                break
    return "\n".join(parts)[:MAX_SUMMARY_LEN]


def has_images(template_path: Path) -> bool:
    """模板中是否存在图片位（用于提示 LLM 标注图片字段）"""
    doc = Document(str(template_path))
    if getattr(doc, "inline_shapes", None) and len(doc.inline_shapes) > 0:
        return True
    for rel in doc.part.rels.values():
        if "image" in str(rel.reltype):
            return True
    return False


# ==================== 文档填充 ====================

def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
    """
    在段落中执行文本替换（支持 {标注符} 和任意文本 key）。
    优先 run 内替换，跨 run 替换仅改命中区间，尽量保留原样式。
    """
    runs = paragraph.runs
    if not runs:
        return False

    norm_repls: list[tuple[str, str]] = []
    for old_text, new_text in (replacements or {}).items():
        old = str(old_text or "")
        if not old:
            continue
        norm_repls.append((old, str(new_text)))
    if not norm_repls:
        return False

    def is_image_run(run) -> bool:
        return bool(
            run._element.findall(qn("w:drawing"))
            or run._element.findall(qn("w:pict"))
        )

    def collect_text_runs():
        text_runs = []
        cursor = 0
        for run in runs:
            if is_image_run(run):
                continue
            txt = run.text or ""
            if not txt:
                continue
            start = cursor
            end = start + len(txt)
            text_runs.append({"run": run, "text": txt, "start": start, "end": end})
            cursor = end
        return text_runs, cursor

    def find_safe_match_start(full_text: str, old_text: str) -> int:
        """
        查找可替换位置：
        - 常规：返回首次命中
        - old 在 new 中（如 old=照片, new={照片}）时，跳过已被 {} 包裹的命中，
          避免反复命中新生成占位符导致死循环。
        """
        if not old_text:
            return -1
        start = 0
        while True:
            idx = full_text.find(old_text, start)
            if idx < 0:
                return -1
            left = full_text[idx - 1] if idx > 0 else ""
            right_idx = idx + len(old_text)
            right = full_text[right_idx] if right_idx < len(full_text) else ""
            if not (left == "{" and right == "}"):
                return idx
            start = idx + len(old_text)

    # 1) 先做 run 内替换（格式保留最佳）
    hit = False
    for run in runs:
        if is_image_run(run):
            continue
        text = run.text or ""
        if not text:
            continue
        replaced = text
        for old, new in norm_repls:
            if old in replaced:
                replaced = replaced.replace(old, new)
        if replaced != text:
            run.text = replaced
            hit = True

    # 2) 跨 run 兜底替换（仅修改命中范围，不清空整段）
    # 关键保护：old/new 自包含时（如 "照片" -> "{照片}"），只替换“未被 {} 包裹”的旧文本，
    # 避免反复命中新生成占位符。
    for old, new in norm_repls:
        while True:
            text_runs, _ = collect_text_runs()
            full_text = "".join(item["text"] for item in text_runs)
            start = find_safe_match_start(full_text, old)
            if start < 0:
                break
            end = start + len(old)
            touched = []
            for item in text_runs:
                if item["end"] <= start or item["start"] >= end:
                    continue
                touched.append(item)
            if not touched:
                break

            first = touched[0]
            last = touched[-1]
            for item in touched:
                run = item["run"]
                txt = run.text or ""
                local_start = max(start, item["start"]) - item["start"]
                local_end = min(end, item["end"]) - item["start"]
                prefix = txt[:local_start]
                suffix = txt[local_end:]
                if item is first and item is last:
                    run.text = prefix + new + suffix
                elif item is first:
                    run.text = prefix + new
                elif item is last:
                    run.text = suffix
                else:
                    run.text = ""
            hit = True

    return hit


def _replace_image_in_paragraph(paragraph, placeholder: str, image_bytes: bytes):
    """
    将段落中的图片标注替换为实际图片。
    找到包含 placeholder 的 run，清除文本并插入图片。
    """
    full_text = _collect_paragraph_text(paragraph)
    if placeholder not in full_text:
        return False

    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = _cleanup_empty_brace_artifacts(
                run.text.replace(placeholder, "")
            )
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(image_bytes)
                tmp_path = tmp.name
            try:
                run.add_picture(tmp_path, width=Inches(1.5))
            finally:
                os.unlink(tmp_path)
            return True

    # 回退：合并后再处理
    runs = paragraph.runs
    if runs:
        runs[0].text = _cleanup_empty_brace_artifacts(
            full_text.replace(placeholder, "")
        )
        for run in runs[1:]:
            run.text = ""
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name
        try:
            runs[0].add_picture(tmp_path, width=Inches(1.5))
        finally:
            os.unlink(tmp_path)
        return True

    return False


def _decode_image_bytes(raw_value: Any) -> bytes | None:
    """兼容 dataURL/base64/url/本地路径 的图片值解码"""
    if raw_value is None:
        return None
    if isinstance(raw_value, bytes):
        return raw_value

    if not isinstance(raw_value, str):
        raw_value = str(raw_value)
    value = raw_value.strip()
    if not value:
        return None

    # data:image/png;base64,xxxx
    if value.startswith("data:image") and "base64," in value:
        try:
            return base64.b64decode(value.split("base64,", 1)[1])
        except Exception:
            return None

    # 纯 base64（常见长度较长）
    if len(value) > 200:
        try:
            return base64.b64decode(value)
        except Exception:
            pass

    # URL 图片
    if value.startswith("http://") or value.startswith("https://"):
        try:
            with urlopen(value, timeout=8) as resp:
                return resp.read()
        except Exception:
            return None

    # 本地路径（仅服务端可访问时）
    p = Path(value)
    if p.exists() and p.is_file():
        try:
            return p.read_bytes()
        except Exception:
            return None
    return None


def _replace_existing_images(doc: Document, image_bytes_list: List[bytes]) -> int:
    """
    当模板没有 {照片} 文本标注时，尝试替换模板中已有图片位。
    典型场景：模板中放了 demo 图片或空白图片框。
    """
    if not image_bytes_list:
        return 0
    image_parts = []
    for rel in doc.part.rels.values():
        if "image" in str(rel.reltype):
            target_part = getattr(rel, "target_part", None)
            if target_part is not None and hasattr(target_part, "_blob"):
                image_parts.append(target_part)
    if not image_parts:
        return 0

    replaced = 0
    for idx, part in enumerate(image_parts):
        if idx >= len(image_bytes_list):
            break
        part._blob = image_bytes_list[idx]
        replaced += 1
    return replaced


def fill_document(
    template_path: Path,
    mappings: List[Dict[str, str]],
    row_data: Dict[str, Any],
) -> bytes:
    """
    根据映射和单行数据填充模板，返回 .docx 字节流。
    
    mappings: [{"placeholder": "{姓名}", "column": "姓名", "type": "text"}, ...]
    row_data: {"姓名": "张三", "照片": "base64...", ...}
    """
    doc = Document(str(template_path))

    # 分离文本替换和图片替换
    text_map: Dict[str, str] = {}
    image_map: Dict[str, bytes] = {}

    for m in mappings:
        col_val = row_data.get(m["column"], "")
        placeholder = _normalize_placeholder(m.get("placeholder"), m.get("column"))
        if not placeholder:
            continue

        if m.get("type") == "image" and col_val:
            decoded = _decode_image_bytes(col_val)
            if decoded:
                image_map[placeholder] = decoded
            else:
                text_map[placeholder] = "(图片不可用)"
        else:
            text_map[placeholder] = _normalize_date_like_text(col_val)

    # 执行替换
    image_hit_map = {k: False for k in image_map.keys()}
    for para in _iter_all_paragraphs(doc):
        _replace_in_paragraph(para, text_map)
        for ph, img_bytes in image_map.items():
            if _replace_image_in_paragraph(para, ph, img_bytes):
                image_hit_map[ph] = True

    # 若模板中没有 {照片} 明文标注，回退替换模板既有图片位
    fallback_images = [
        image_map[ph] for ph, hit in image_hit_map.items() if not hit
    ]
    if fallback_images:
        _replace_existing_images(doc, fallback_images)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==================== 高保真 docx -> HTML ====================

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def _run_to_html(run, fallback_part=None) -> str:
    """将单个 Run 转为带 inline style 的 <span>"""
    text = run.text or ""
    drawing_els = run._element.findall(qn("w:drawing"))
    pict_els = run._element.findall(qn("w:pict"))
    image_html = ""
    if drawing_els or pict_els:
        image_html = _extract_images_from_run(run, drawing_els, pict_els, fallback_part)
    if not text and not image_html:
        return ""

    parts: List[str] = []
    font = run.font

    if font.bold:
        parts.append("font-weight:bold")
    if font.italic:
        parts.append("font-style:italic")
    if font.underline:
        parts.append("text-decoration:underline")

    if font.size:
        pt = font.size.pt
        parts.append(f"font-size:{pt}pt")
    if font.name:
        parts.append(f"font-family:{escape(font.name)}")
    if font.color and font.color.rgb:
        parts.append(f"color:#{font.color.rgb}")

    text_html = ""
    if text:
        escaped = escape(text)
        text_html = escaped if not parts else f'<span style="{";".join(parts)}">{escaped}</span>'
    return text_html + image_html


def _extract_images_from_run(run, drawing_els, pict_els, fallback_part=None) -> str:
    """从 w:drawing / w:pict 中提取图片并转为 <img> 标签"""
    part = fallback_part
    if part is None:
        try:
            part = run.part
        except Exception:
            part = None
    if part is None:
        return ""

    results = []
    # 1) DrawingML: a:blip r:embed
    for drawing in drawing_els:
        blip_els = drawing.findall(
            ".//" + qn("a:blip")
        )
        for blip in blip_els:
            embed_id = blip.get(qn("r:embed"))
            if not embed_id:
                continue
            rel = part.rels.get(embed_id)
            if not rel:
                continue
            target_part = getattr(rel, "target_part", None)
            if target_part is None or not hasattr(target_part, "_blob"):
                continue
            blob = target_part._blob
            content_type = getattr(target_part, "content_type", "image/png")
            b64 = base64.b64encode(blob).decode("ascii")
            results.append(
                f'<img src="data:{content_type};base64,{b64}" '
                f'data-embed-id="{embed_id}" '
                f'style="max-width:100%;height:auto;vertical-align:middle">'
            )
    # 2) VML: v:imagedata r:id
    for pict in pict_els:
        image_data_els = pict.findall(".//{urn:schemas-microsoft-com:vml}imagedata")
        for image_data in image_data_els:
            rid = image_data.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if not rid:
                continue
            rel = part.rels.get(rid)
            if not rel:
                continue
            target_part = getattr(rel, "target_part", None)
            if target_part is None or not hasattr(target_part, "_blob"):
                continue
            blob = target_part._blob
            content_type = getattr(target_part, "content_type", "image/png")
            b64 = base64.b64encode(blob).decode("ascii")
            results.append(
                f'<img src="data:{content_type};base64,{b64}" '
                f'data-embed-id="{rid}" '
                f'style="max-width:100%;height:auto;vertical-align:middle">'
            )
    return "".join(results)


def _paragraph_to_html(para, fallback_part=None) -> str:
    """将段落转为 <p> 标签，保留对齐和 run 格式"""
    style_parts: List[str] = []
    align = para.alignment
    if align is not None and align in _ALIGN_MAP:
        style_parts.append(f"text-align:{_ALIGN_MAP[align]}")

    # 段落间距
    pf = para.paragraph_format
    if pf.space_before and pf.space_before.pt:
        style_parts.append(f"margin-top:{pf.space_before.pt}pt")
    if pf.space_after and pf.space_after.pt:
        style_parts.append(f"margin-bottom:{pf.space_after.pt}pt")
    if pf.first_line_indent and pf.first_line_indent.pt:
        style_parts.append(f"text-indent:{pf.first_line_indent.pt}pt")

    runs_html = "".join(_run_to_html(r, fallback_part) for r in para.runs)

    # 空段落也要保留结构（Word 中的空行）
    if not runs_html.strip():
        runs_html = "&nbsp;"

    style_attr = f' style="{";".join(style_parts)}"' if style_parts else ""
    return f"<p{style_attr}>{runs_html}</p>"


def _table_to_html(table, doc_part=None) -> str:
    """将表格转为 <table>，处理合并单元格 (gridSpan / vMerge)"""
    tbl = table._tbl

    # 第一遍：收集每行 tc 元素及合并信息，并计算逻辑网格列范围
    # 注意：不能用“数组下标”判断 vMerge 对齐，因为 gridSpan 会改变列位。
    rows_data: List[List[Dict]] = []
    for tr in tbl.findall(qn("w:tr")):
        row_cells = []
        grid_cursor = 0
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            colspan = 1
            vmerge_status = "none"

            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    colspan = int(gs.get(qn("w:val"), "1"))
                vm = tcPr.find(qn("w:vMerge"))
                if vm is not None:
                    val = vm.get(qn("w:val"), "")
                    vmerge_status = "restart" if val == "restart" else "continue"

            row_cells.append({
                "tc": tc,
                "colspan": colspan,
                "vmerge": vmerge_status,
                "grid_start": grid_cursor,
                "grid_end": grid_cursor + colspan,
            })
            grid_cursor += colspan
        rows_data.append(row_cells)

    # 第二遍：计算 rowspan（按逻辑网格列向下扫描 "continue" 单元格）
    # 这样可正确处理“合并列 + 纵向合并”组合场景，避免预览错位。
    for row in rows_data:
        for cell_info in row:
            if cell_info["vmerge"] == "continue":
                cell_info["skip"] = True

    for ri, row in enumerate(rows_data):
        for cell_info in row:
            if cell_info["vmerge"] == "restart":
                span = 1
                start_col = cell_info["grid_start"]
                end_col = cell_info["grid_end"]
                for ri2 in range(ri + 1, len(rows_data)):
                    next_row = rows_data[ri2]
                    matched_continue = next(
                        (
                            c for c in next_row
                            if c["grid_start"] == start_col
                            and c["grid_end"] == end_col
                            and c["vmerge"] == "continue"
                        ),
                        None,
                    )
                    if matched_continue is None:
                        break
                    span += 1
                cell_info["rowspan"] = span

    # 第三遍：渲染 HTML
    parts = [
        '<table border="1" cellspacing="0" cellpadding="6" '
        'style="border-collapse:collapse;width:100%;margin:8px 0;border:1px solid #222">'
    ]
    for row in rows_data:
        parts.append("<tr>")
        for cell_info in row:
            if cell_info.get("skip"):
                continue
            attrs: List[str] = []
            if cell_info["colspan"] > 1:
                attrs.append(f'colspan="{cell_info["colspan"]}"')
            rowspan = cell_info.get("rowspan", 1)
            if rowspan > 1:
                attrs.append(f'rowspan="{rowspan}"')

            # 单元格样式
            tc = cell_info["tc"]
            td_style = _extract_cell_style(tc)
            # 兜底单元格边框：避免浏览器/样式重置导致预览无网格线
            if td_style:
                if "border" not in td_style:
                    td_style = f"{td_style};border:1px solid #222"
                attrs.append(f'style="{td_style}"')
            else:
                attrs.append('style="border:1px solid #222"')

            attr_str = (" " + " ".join(attrs)) if attrs else ""

            # 单元格内容：遍历段落
            cell_html = []
            for p_el in tc.findall(qn("w:p")):
                from docx.text.paragraph import Paragraph
                p_obj = Paragraph(p_el, tc)
                cell_html.append(_paragraph_to_html(p_obj, doc_part))
            content = "".join(cell_html) or "&nbsp;"
            parts.append(f"<td{attr_str}>{content}</td>")
        parts.append("</tr>")
    parts.append("</table>")
    return "".join(parts)


def _extract_cell_style(tc) -> str:
    """从 tc XML 提取单元格样式（背景色、宽度、垂直对齐）"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return ""
    parts: List[str] = []
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"), "")
        if fill and fill.lower() not in ("auto", "ffffff"):
            parts.append(f"background-color:#{fill}")
    tcw = tcPr.find(qn("w:tcW"))
    if tcw is not None:
        w = tcw.get(qn("w:w"))
        typ = tcw.get(qn("w:type"), "")
        if w and typ == "dxa":
            try:
                # dxa: 1/20 point，转为 pt 保留 Word 单元格宽度信息
                parts.append(f"width:{int(w) / 20:.2f}pt")
            except Exception:
                pass
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is not None:
        val = vAlign.get(qn("w:val"), "")
        if val == "center":
            parts.append("vertical-align:middle")
        elif val == "bottom":
            parts.append("vertical-align:bottom")
    return ";".join(parts)


def doc_to_html_hifi(doc: Document) -> str:
    """
    高保真 docx -> HTML。
    按文档 body 子元素顺序遍历，保留段落/表格交替顺序、
    run 级格式、合并单元格、内嵌图片。
    """
    body = doc.element.body
    blocks: List[str] = []

    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, body)
            blocks.append(_paragraph_to_html(para))
        elif tag == "tbl":
            from docx.table import Table
            tbl = Table(child, body)
            blocks.append(_table_to_html(tbl, doc.part))

    if not blocks:
        blocks.append("<p>（文档暂无可预览文本）</p>")

    return (
        '<div style="background:#fff;color:#222;padding:24px 28px;'
        'font-family:SimSun,宋体,serif;line-height:1.7;max-width:800px;margin:0 auto">'
        + "".join(blocks)
        + "</div>"
    )


def build_preview_html(
    template_path: Path,
    mappings: List[Dict[str, Any]],
    row_data: Dict[str, Any],
    mode: str = "filled",
) -> str:
    """按当前映射+样本行生成高保真在线预览 HTML"""
    if mode == "annotated":
        doc = Document(str(template_path))
        return doc_to_html_hifi(doc)

    filled = fill_document(template_path, mappings, row_data)
    doc = Document(io.BytesIO(filled))
    return doc_to_html_hifi(doc)


# ==================== 文档标注（修改 docx 本体） ====================

def annotate_doc_with_replacements(
    template_path: Path,
    replacements: Dict[str, str],
) -> str:
    """
    在 docx 中执行文本替换（demo数据 -> {字段}），覆盖保存并返回高保真 HTML。
    replacements: {"张三": "{姓名}", "2024年1月": "{日期}", ...}
    """
    doc = Document(str(template_path))
    for para in _iter_all_paragraphs(doc):
        _replace_in_paragraph(para, replacements)
    doc.save(str(template_path))
    doc = Document(str(template_path))
    return doc_to_html_hifi(doc)


def normalize_placeholder_artifacts_in_doc(
    template_path: Path,
    excel_columns: List[str],
) -> str:
    """
    清洗模板中的占位符噪音：
    - {{{列名}}} / {{列名}} -> {列名}
    - {{{}}} / {{}} 等空占位符 -> 删除
    覆盖保存并返回新的高保真 HTML。
    """
    doc = Document(str(template_path))
    replacements: Dict[str, str] = {}

    for col in excel_columns or []:
        canonical = _normalize_placeholder(col, col)
        if not canonical:
            continue
        name = canonical[1:-1]
        for level in range(2, 7):
            replacements["{" * level + name + "}" * level] = canonical
            replacements["{" * level + f" {name} " + "}" * level] = canonical

    # 清理空占位符噪音
    for noise in ("{{}}", "{{{}}}", "{{{{}}}}", "{ }", "{{ }}", "{{{ }}}"):
        replacements[noise] = ""

    if replacements:
        for para in _iter_all_paragraphs(doc):
            _replace_in_paragraph(para, replacements)

    doc.save(str(template_path))
    doc = Document(str(template_path))
    return doc_to_html_hifi(doc)


def manual_annotate_in_doc(
    template_path: Path,
    original_text: str,
    field_name: str,
) -> str:
    """
    单次手工标注：将 original_text 替换为 {field_name}，
    覆盖保存 docx 并返回新 HTML。
    """
    placeholder = _normalize_placeholder(field_name, field_name)
    if not placeholder:
        placeholder = "{字段}"
    return annotate_doc_with_replacements(
        template_path, {original_text: placeholder}
    )


def manual_annotate_image_in_doc(
    template_path: Path,
    field_name: str,
    embed_id: str | None = None,
) -> str:
    """
    手工图片标注：将匹配图片位替换为 {field_name} 文本占位符，
    覆盖保存 docx 并返回新 HTML。
    """
    doc = Document(str(template_path))
    placeholder = _normalize_placeholder(field_name, field_name)
    if not placeholder:
        placeholder = "{图片}"

    def _replace_para_image(para, target_embed_id: str | None) -> bool:
        target_run = None
        for run in para.runs:
            drawings = run._element.findall(qn("w:drawing"))
            if not drawings:
                continue
            if target_embed_id:
                hit = False
                for d in drawings:
                    for blip in d.findall(".//" + qn("a:blip")):
                        if blip.get(qn("r:embed")) == target_embed_id:
                            hit = True
                            break
                    if hit:
                        break
                if not hit:
                    continue
            run.clear()
            run.text = placeholder
            target_run = run
            break
        if target_run is None:
            return False

        # 清理同段落中可能存在的多层花括号残留，避免出现 {{{相片}}}
        text_runs = [r for r in para.runs if not r._element.findall(qn("w:drawing")) and not r._element.findall(qn("w:pict"))]
        if text_runs:
            merged = "".join(r.text or "" for r in text_runs)
            normalized = _normalize_nested_placeholder_braces(merged)
            if normalized != merged:
                text_runs[0].text = normalized
                for r in text_runs[1:]:
                    r.text = ""
        return True

    replaced = False
    for para in _iter_all_paragraphs(doc):
        if _replace_para_image(para, embed_id):
            replaced = True
            break
    if not replaced and embed_id:
        for para in _iter_all_paragraphs(doc):
            if _replace_para_image(para, None):
                replaced = True
                break
    if not replaced:
        for para in doc.paragraphs:
            if not para.text.strip():
                para.text = placeholder
                replaced = True
                break
    if not replaced:
        doc.add_paragraph(placeholder)

    doc.save(str(template_path))
    doc = Document(str(template_path))
    return doc_to_html_hifi(doc)


# ==================== 文件命名 ====================

def resolve_filename(
    pattern: str,
    row_data: Dict[str, Any],
    index: int,
    image_columns: set[str] | None = None,
) -> str:
    """
    解析文件名模式，替换 {列名} 占位符。
    内置变量 {_index} 代表序号（从 1 开始）。
    """
    name = pattern.replace("{_index}", str(index))
    image_columns = image_columns or set()
    for key, val in row_data.items():
        if key in image_columns:
            name = name.replace("{" + key + "}", "")
            continue
        normalized = _normalize_date_like_text(val).replace("/", "-")
        name = name.replace("{" + key + "}", normalized)
    # 清除文件名中的非法字符
    name = re.sub(r'[\\/:*?"<>|]', '_', name)
    return name.strip() or f"文档_{index}"


# ==================== 批量生成 + ZIP 打包 ====================

def batch_generate(
    template_path: Path,
    mappings: List[Dict[str, str]],
    rows: List[Dict[str, Any]],
    filename_pattern: str,
) -> Tuple[str, Path]:
    """
    批量生成 Word 文档并打包 ZIP，返回 (task_id, zip_path)。
    """
    date_dir = _today_dir()
    task_id = uuid.uuid4().hex[:12]
    zip_path = date_dir / f"{task_id}.zip"

    image_columns = {
        str(m.get("column") or "").strip()
        for m in mappings
        if str(m.get("type") or "text").strip() == "image"
    }

    with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
        for idx, row in enumerate(rows, start=1):
            doc_bytes = fill_document(template_path, mappings, row)
            fname = resolve_filename(filename_pattern, row, idx, image_columns=image_columns) + ".docx"
            zf.writestr(fname, doc_bytes)

    logger.info("批量生成完成: task=%s count=%d", task_id, len(rows))
    return task_id, zip_path
